"""DOCX writer from Document IR with font mapping and page-break options."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document as WordDocument
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.text import WD_BREAK
from docx.shared import Pt

from ir import Document, Page, ParagraphBlock, Run, TableBlock


@dataclass(slots=True, frozen=True)
class FontMapRule:
    contains: str
    word_font: str


DEFAULT_FONT = "Calibri"

DEFAULT_FONT_RULES = (
    FontMapRule("thsarabun", "TH Sarabun New"),
    FontMapRule("angsana", "Angsana New"),
    FontMapRule("leelawadee", "Leelawadee UI"),
    FontMapRule("times", "Times New Roman"),
    FontMapRule("helvetica", "Arial"),
    FontMapRule("arial", "Arial"),
    FontMapRule("courier", "Courier New"),
)


class DocxExporter:
    def __init__(
        self,
        *,
        include_page_breaks: bool = True,
        default_font: str = DEFAULT_FONT,
        font_rules: tuple[FontMapRule, ...] = DEFAULT_FONT_RULES,
    ) -> None:
        self._include_page_breaks = include_page_breaks
        self._default_font = default_font
        self._font_rules = font_rules

    def export(self, document: Document, output_path: str | Path) -> None:
        word_doc = WordDocument()
        if document.pages:
            self._apply_page_layout(word_doc.sections[0], document.pages[0])
        for page_idx, page in enumerate(document.pages):
            if self._include_page_breaks and page_idx > 0:
                section = word_doc.add_section(WD_SECTION_START.NEW_PAGE)
                self._apply_page_layout(section, page)
            for block in page.blocks:
                if isinstance(block, ParagraphBlock):
                    self._write_paragraph(word_doc, block.runs)
                elif isinstance(block, TableBlock):
                    self._write_table(word_doc, block)
            if not self._include_page_breaks and page_idx < len(document.pages) - 1:
                breaker = word_doc.add_paragraph()
                breaker.add_run().add_break(WD_BREAK.PAGE)
        word_doc.save(str(output_path))

    def _write_paragraph(self, word_doc: WordDocument, runs: list[Run]) -> None:
        paragraph = word_doc.add_paragraph()
        for run in runs:
            self._apply_run(paragraph.add_run(run.text), run)

    def _write_table(self, word_doc: WordDocument, table_block: TableBlock) -> None:
        rows = table_block.rows
        if not rows:
            return
        max_cols = max(len(row) for row in rows)
        table = word_doc.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx in range(max_cols):
                cell = table.cell(r_idx, c_idx)
                paragraph = cell.paragraphs[0]
                if c_idx >= len(row):
                    paragraph.add_run("")
                    continue
                for run in row[c_idx].runs:
                    self._apply_run(paragraph.add_run(run.text), run)

    def _map_font_name(self, pdf_font_name: str | None) -> str:
        if not pdf_font_name:
            return self._default_font
        lower = pdf_font_name.lower()
        for rule in self._font_rules:
            if rule.contains in lower:
                return rule.word_font
        return self._default_font

    def _apply_run(self, word_run, ir_run: Run) -> None:
        word_run.bold = ir_run.bold
        word_run.italic = ir_run.italic
        word_run.font.name = self._map_font_name(ir_run.font_name)
        if ir_run.size_pt is not None:
            word_run.font.size = Pt(ir_run.size_pt)

    def _apply_page_layout(self, section, page: Page) -> None:
        width = page.width_pt
        height = page.height_pt
        if width is None or height is None:
            return
        section.page_width = Pt(width)
        section.page_height = Pt(height)
        if width > height:
            section.orientation = WD_ORIENT.LANDSCAPE
        else:
            section.orientation = WD_ORIENT.PORTRAIT


def export_docx(
    document: Document,
    output_path: str | Path,
    *,
    include_page_breaks: bool = True,
) -> None:
    DocxExporter(include_page_breaks=include_page_breaks).export(document, output_path)
